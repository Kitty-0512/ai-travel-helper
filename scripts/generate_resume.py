"""Generate a well-formatted Chinese resume in Word (.docx)."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "简历_AI应用开发实习生.docx")

# ── colour palette ──────────────────────────────────────────
PRIMARY   = RGBColor(0x1A, 0x3A, 0x5C)   # deep navy
ACCENT    = RGBColor(0x2B, 0x6C, 0xB0)   # medium blue
DARK      = RGBColor(0x1E, 0x1E, 0x1E)   # almost black
GREY      = RGBColor(0x66, 0x66, 0x66)   # body grey
LIGHT_BG  = RGBColor(0xF0, 0xF4, 0xF8)   # section bg tint
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LINE_CLR  = RGBColor(0xCC, 0xD5, 0xE0)   # subtle line

# ── helpers ─────────────────────────────────────────────────


def set_font(run, name_cn="微软雅黑", name_en="Calibri", size=Pt(10), color=DARK, bold=False):
    """Set both Chinese and Latin font on a run."""
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name_cn)
    rFonts.set(qn("w:ascii"), name_en)
    rFonts.set(qn("w:hAnsi"), name_en)


def add_run(paragraph, text, name_cn="微软雅黑", name_en="Calibri",
            size=Pt(10), color=DARK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_font(run, name_cn, name_en, size, color, bold)
    if italic:
        run.italic = True
    return run


def set_paragraph_spacing(paragraph, before=Pt(0), after=Pt(0), line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = before
    pf.space_after = after
    pf.line_spacing = line


def add_section_header(doc, title):
    """Thin coloured bar + section title."""
    # coloured line
    p_line = doc.add_paragraph()
    set_paragraph_spacing(p_line, before=Pt(14), after=Pt(2))
    pPr = p_line._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2B6CB0")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # title
    p_title = doc.add_paragraph()
    set_paragraph_spacing(p_title, before=Pt(0), after=Pt(4), line=1.0)
    add_run(p_title, title, size=Pt(13), color=PRIMARY, bold=True)
    return p_title


def add_job_header(doc, company, role, dates, right_text=""):
    """Company | Role ········· Dates"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(8), after=Pt(1), line=1.15)
    add_run(p, company, size=Pt(10.5), color=DARK, bold=True)
    add_run(p, " ｜ ", size=Pt(10.5), color=GREY)
    add_run(p, role, size=Pt(10.5), color=DARK)
    if dates:
        add_run(p, " ｜ ", size=Pt(10.5), color=GREY)
        add_run(p, dates, size=Pt(9.5), color=GREY)
    return p


def add_bullet(doc, text_parts, indent=Cm(0.6)):
    """Add a bullet point. text_parts is a list of (text, bold, color) tuples."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(1), after=Pt(1), line=1.2)
    pf = p.paragraph_format
    pf.left_indent = indent
    pf.first_line_indent = Cm(-0.35)

    add_run(p, "• ", size=Pt(9.5), color=ACCENT, bold=True)
    for part in text_parts:
        text, bold, color = part[0], part[1] if len(part) > 1 else False, part[2] if len(part) > 2 else DARK
        add_run(p, text, size=Pt(9.5), color=color, bold=bold)
    return p


def add_tech_stack(doc, stack, indent=Cm(0.6)):
    """Smaller grey line showing tech stack."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(0), after=Pt(2), line=1.15)
    pf = p.paragraph_format
    pf.left_indent = indent
    add_run(p, stack, size=Pt(8.5), color=GREY, name_en="Consolas")
    return p


def add_proj_title(doc, name, stack, right_text=""):
    """Project name with tech stack below."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=Pt(10), after=Pt(1), line=1.15)
    add_run(p, f"▸ {name}", size=Pt(10.5), color=DARK, bold=True)
    if right_text:
        add_run(p, "  ｜ ", size=Pt(9.5), color=GREY)
        add_run(p, right_text, size=Pt(9.5), color=GREY)
    return p


# ============================================================
# BUILD
# ============================================================

doc = Document()

# ── page setup ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

# set default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

# ── HEADER ──────────────────────────────────────────────────
p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p_name, before=Pt(0), after=Pt(2), line=1.0)
add_run(p_name, "张同学", size=Pt(22), color=PRIMARY, bold=True)

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p_info, before=Pt(0), after=Pt(0), line=1.2)
for text in ["📧 zhang@example.com", " ｜ ", "📱 138-xxxx-xxxx", " ｜ ",
             "🌐 github.com/zhang", " ｜ ", "📍 上海"]:
    bold = text.startswith("📧") or text.startswith("📱") or text.startswith("🌐") or text.startswith("📍")
    add_run(p_info, text, size=Pt(9), color=GREY, bold=bold)

# thin divider under header
p_div = doc.add_paragraph()
set_paragraph_spacing(p_div, before=Pt(6), after=Pt(0))
pPr = p_div._element.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "6")
bottom.set(qn("w:color"), "2B6CB0")
pBdr.append(bottom)
pPr.append(pBdr)

# white space
set_paragraph_spacing(doc.add_paragraph(), before=Pt(0), after=Pt(0), line=0.2)

# ── EDUCATION ───────────────────────────────────────────────
add_section_header(doc, "教育背景")

p_edu = doc.add_paragraph()
set_paragraph_spacing(p_edu, before=Pt(4), after=Pt(2), line=1.15)
add_run(p_edu, "XX 大学", size=Pt(10.5), color=DARK, bold=True)
add_run(p_edu, " ｜ 计算机科学与技术 本科 ｜ 2026 年毕业", size=Pt(9.5), color=GREY)

# ── INTERNSHIP ──────────────────────────────────────────────
add_section_header(doc, "实习经历")

add_job_header(doc, "上海云璨信息技术有限公司", "AI 应用开发实习生", "2025.05 — 至今")

add_bullet(doc, [
    ("参与企业 AI 智能分析平台建设，基于 ",),
    ("FastAPI + LangChain + DeepSeek", True),
    (" 实现自然语言数据分析与日志诊断，设计 Agent 路由、工具调用与 SSE 流式交互",),
])

add_bullet(doc, [
    ("开发多站点 SEO 数据分析平台，对接 ",),
    ("百度统计、Bing、Google Search Console", True),
    (" 等 API，实现数据采集、Token 管理与可视化分析",),
])

add_bullet(doc, [
    ("基于 ",),
    ("MCP SDK", True),
    (" 封装企业业务工具，调研 ",),
    ("RAG 知识库", True),
    (" 与 FastGPT 完成模型能力扩展；使用 ",),
    ("Docker Compose + Nginx", True),
    (" 完成服务部署与维护",),
])

# ── PROJECTS ────────────────────────────────────────────────
add_section_header(doc, "项目经历")

# ---- Project 1 ----
add_proj_title(doc, "AI 统一分析网关", "Vue3 · FastAPI · LangChain · DeepSeek · SQLite · Docker", "全栈独立开发")

add_bullet(doc, [
    ("设计 ",),
    ("双场景统一路由", True),
    ("：结合文件类型、关键词与 LLM 分类将请求分流至 Text-to-SQL 或日志诊断链路，共用 SSE 流式交互",),
])

add_bullet(doc, [
    ("实现 ",),
    ("Text-to-SQL 闭环", True),
    ("：Excel/CSV 入库 → Schema 组装 Prompt → LLM 生成 SQL → 安全校验与字段存在性校验 → 执行；失败后结合错误信息与 Schema ",),
    ("自动修复（最多 3 轮）", True),
])

add_bullet(doc, [
    ("实现日志诊断链路：支持 Docker / Nginx / 应用堆栈日志分析，输出 ",),
    ("问题原因、关键证据与排查建议", True),
    ("，首轮完整报告 + 追问精简回答两种模式",),
])

add_bullet(doc, [
    ("前端基于 ",),
    ("ECharts", True),
    (" 可视化查询结果，SSE 流式展示模型输出与工具执行过程，支持 PDF 报告导出",),
])

# ---- Project 2 ----
add_proj_title(doc, "AI 旅行规划助手", "Vue3 · TypeScript · FastAPI · DeepSeek · MCP SDK · PostgreSQL · pgvector · Redis · Docker", "全栈独立开发")

add_bullet(doc, [
    ("设计 ",),
    ("Planner → Executor → Finalize 三阶段 Agent 架构", True),
    ("：Planner 拆解需求为 JSON 任务计划并支持规则兜底，Executor 串行调用 MCP 工具获取真实数据，Finalize 生成结构化行程并通过 ",),
    ("7 种 SSE 事件", True),
    (" 流式推送",),
])

add_bullet(doc, [
    ("基于 ",),
    ("MCP Python SDK", True),
    (" 自建 stdio 传输工具服务，封装高德地图天气/POI/路线 API；修复 ",),
    ("PEP 563 延迟注解与 SDK 类型检查冲突", True),
    (" 的兼容性问题",),
])

add_bullet(doc, [
    ("实现 ",),
    ("多轮对话行程合并", True),
    ("：ReAct 循环支持追问调整，Session 持久化 PostgreSQL；通过注入已有行程 JSON 解决上下文截断导致旧数据丢失的问题",),
])

add_bullet(doc, [
    ("构建 ",),
    ("长期偏好记忆", True),
    ("：LLM 提取用户偏好（6 维度）→ PostgreSQL + pgvector（HNSW 索引）→ 后续对话注入 Prompt 实现个性化",),
])

add_bullet(doc, [
    ("前端实现 ",),
    ("Haversine 贪心最近邻路径优化", True),
    (" 与高德地图多日分组展示，支持 PDF 行程导出",),
])

# ── SKILLS ──────────────────────────────────────────────────
add_section_header(doc, "专业技能")

p_skills = doc.add_paragraph()
set_paragraph_spacing(p_skills, before=Pt(4), after=Pt(2), line=1.3)
pf = p_skills.paragraph_format
pf.left_indent = Cm(0.6)
pf.first_line_indent = Cm(-0.35)

skills = [
    ("语言：", "Python · TypeScript · SQL", True),
    ("框架：", "FastAPI · Vue3 · LangChain · Tailwind CSS", False),
    ("LLM/AI：", "DeepSeek · MCP SDK · RAG · pgvector · Agent 编排", False),
    ("数据：", "PostgreSQL · Redis · SQLite · ECharts", False),
    ("工程：", "Docker · Nginx · Git · RESTful API · SSE 流式协议", False),
]

for i, (label, content, first) in enumerate(skills):
    if i > 0:
        add_run(p_skills, "  │  ", size=Pt(9), color=LINE_CLR)
    add_run(p_skills, label, size=Pt(9.5), color=DARK, bold=True)
    add_run(p_skills, content, size=Pt(9.5), color=DARK)

# ── SAVE ────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Resume saved to: {OUTPUT}")
